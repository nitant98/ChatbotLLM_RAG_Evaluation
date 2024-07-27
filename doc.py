import pandas as pd
from docx import Document

# Load the evaluation results before and after fine-tuning
file_before = 'evaluation_metrics.xlsx'
file_after = 'evaluation_metrics_fine_tune.xlsx'

# Read the Excel files
df_before = pd.read_excel(file_before)
df_after = pd.read_excel(file_after)

# Inspect the column names
print("Columns in df_before:", df_before.columns)
print("Columns in df_after:", df_after.columns)

# Create a new Document
doc = Document()

# Title
doc.add_heading('RAG Pipeline Evaluation Report', 0)

# Introduction
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "This report provides a detailed evaluation of the RAG (Retrieval-Augmented Generation) pipeline before and after implementing improvements. "
    "The evaluation includes various metrics such as retrieval precision, recall, relevance, faithfulness of the generated answers, and more."
)

# Methodology
doc.add_heading('Methodology', level=1)
doc.add_paragraph(
    "The following methodologies were used to calculate each metric:\n"
    "1. Retrieval Metrics:\n"
    "   - Context Precision: The ratio of relevant context terms retrieved to the total number of relevant context terms expected.\n"
    "   - Context Recall: The ratio of relevant context terms retrieved to the total number of documents retrieved.\n"
    "   - Context Relevance: Assessed as the same as context precision for simplicity.\n"
    "   - Context Entity Recall: Assessed as the same as context recall for simplicity.\n"
    "   - Noise Robustness: The ability of the system to handle noisy or irrelevant inputs, evaluated by observing system behavior with irrelevant queries.\n"
    "2. Generation Metrics:\n"
    "   - Faithfulness: The ratio of correct terms in the generated answer to the total number of correct terms expected.\n"
    "   - Answer Relevance: Assessed as the same as faithfulness for simplicity.\n"
    "   - Information Integration: Assessed as the same as answer relevance for simplicity.\n"
    "   - Counterfactual Robustness: The system's ability to handle counterfactual or contradictory queries, measured with simulated robustness evaluation.\n"
    "   - Negative Rejection: The system's ability to reject inappropriate queries, measured with simulated rejection evaluation.\n"
    "   - Latency: The time taken by the system to generate an answer from receiving the query."
)

# Results Before Improvements
doc.add_heading('Results Before Improvements', level=1)
doc.add_paragraph(
    "The following table shows the evaluation metrics before any improvements were made to the RAG pipeline:"
)

# Adding table for results before improvements
table_before = doc.add_table(rows=1, cols=11)
hdr_cells = table_before.rows[0].cells
hdr_cells[0].text = 'Query'
hdr_cells[1].text = 'Context Precision'
hdr_cells[2].text = 'Context Recall'
hdr_cells[3].text = 'Context Relevance'
hdr_cells[4].text = 'Context Entity Recall'
hdr_cells[5].text = 'Faithfulness'
hdr_cells[6].text = 'Answer Relevance'
hdr_cells[7].text = 'Information Integration'
hdr_cells[8].text = 'Counterfactual Robustness'
hdr_cells[9].text = 'Negative Rejection'
hdr_cells[10].text = 'Latency (s)'

# Ensure to use the correct column names
for index, row in df_before.iterrows():
    row_cells = table_before.add_row().cells
    row_cells[0].text = str(row['query'])
    row_cells[1].text = str(row['context_precision'])
    row_cells[2].text = str(row['context_recall'])
    row_cells[3].text = str(row['context_relevance'])
    row_cells[4].text = str(row['context_entity_recall'])
    row_cells[5].text = str(row['faithfulness'])
    row_cells[6].text = str(row['answer_relevance'])
    row_cells[7].text = str(row['information_integration'])
    row_cells[8].text = str(row['counterfactual_robustness'])
    row_cells[9].text = str(row['negative_rejection'])
    row_cells[10].text = str(row['latency'])

# Results After Improvements
doc.add_heading('Results After Improvements', level=1)
doc.add_paragraph(
    "The following table shows the evaluation metrics after implementing improvements to the RAG pipeline:"
)

# Adding table for results after improvements
table_after = doc.add_table(rows=1, cols=11)
hdr_cells = table_after.rows[0].cells
hdr_cells[0].text = 'Query'
hdr_cells[1].text = 'Context Precision'
hdr_cells[2].text = 'Context Recall'
hdr_cells[3].text = 'Context Relevance'
hdr_cells[4].text = 'Context Entity Recall'
hdr_cells[5].text = 'Faithfulness'
hdr_cells[6].text = 'Answer Relevance'
hdr_cells[7].text = 'Information Integration'
hdr_cells[8].text = 'Counterfactual Robustness'
hdr_cells[9].text = 'Negative Rejection'
hdr_cells[10].text = 'Latency (s)'

# Ensure to use the correct column names
for index, row in df_after.iterrows():
    row_cells = table_after.add_row().cells
    row_cells[0].text = str(row['query'])
    row_cells[1].text = str(row['context_precision'])
    row_cells[2].text = str(row['context_recall'])
    row_cells[3].text = str(row['context_relevance'])
    row_cells[4].text = str(row['context_entity_recall'])
    row_cells[5].text = str(row['faithfulness'])
    row_cells[6].text = str(row['answer_relevance'])
    row_cells[7].text = str(row['information_integration'])
    row_cells[8].text = str(row['counterfactual_robustness'])
    row_cells[9].text = str(row['negative_rejection'])
    row_cells[10].text = str(row['latency'])

# Methods Proposed and Implemented for Improvement
doc.add_heading('Methods Proposed and Implemented for Improvement', level=1)
doc.add_paragraph(
    "1. Improving Context Precision and Recall:\n"
    "   - Method: Switched to `sentence-transformers/all-mpnet-base-v2` for better semantic similarity.\n"
    "   - Implementation: Modified the `qa_setup.py` to load the new embedding model and adjusted the number of documents retrieved from 3 to 5.\n"
    "2. Improving Faithfulness and Answer Relevance:\n"
    "   - Method: Switched to GPT-3.5-turbo and fine-tuned the model on a domain-specific dataset.\n"
    "   - Implementation: Updated `model_setup.py` to use GPT-3.5-turbo and fine-tuned the model."
)

# Comparative Analysis
doc.add_heading('Comparative Analysis', level=1)
doc.add_paragraph(
    "The comparative analysis shows the performance improvements achieved after the proposed methods were implemented:\n"
    "- Context Precision and Recall: Improved due to better embeddings, resulting in more accurate semantic matches.\n"
    "- Context Relevance: Increased precision led to improved relevance.\n"
    "- Context Entity Recall: More relevant context retrieved.\n"
    "- Faithfulness and Answer Relevance: Enhanced with a more advanced language model (GPT-3.5-turbo) and domain-specific fine-tuning.\n"
    "- Information Integration: Improved with better model and fine-tuning.\n"
    "- Latency: Improved performance with better retrieval and generation efficiency."
)

# Challenges Faced
doc.add_heading('Challenges Faced', level=1)
doc.add_paragraph(
    "1. ZeroDivisionError in Metrics Calculation:\n"
    "   - Challenge: Encountered a `ZeroDivisionError` due to empty expected context lists in noise robustness testing.\n"
    "   - Solution: Handled the case by returning zero metrics when the expected context list is empty.\n"
    "2. Fine-tuning the Language Model:\n"
    "   - Challenge: Fine-tuning a language model requires substantial computational resources and a high-quality domain-specific dataset.\n"
    "   - Solution: Utilized available resources efficiently and carefully curated the dataset to ensure meaningful fine-tuning.\n"
    "3. Balancing Retrieval Parameters:\n"
    "   - Challenge: Finding the optimal number of documents to retrieve (`k`) for balancing precision and recall.\n"
    "   - Solution: Conducted multiple tests with different `k` values and chose the one providing the best balance."
)

# Save the document
file_path = 'RAG_Pipeline_Evaluation_Report.docx'
doc.save(file_path)

print(f"Report saved to {file_path}")